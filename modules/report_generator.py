import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# -------------------------------------------------------------------
# Generic helpers (fonts, paragraphs, images, tables)
# -------------------------------------------------------------------

def set_paragraph_font(paragraph, size=10):
    """Apply a uniform font size to all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.size = Pt(size)


def insert_section_title(doc, title, level=1):
    """Insert a section title using the template heading style."""
    heading = doc.add_heading(title, level=level)
    set_paragraph_font(heading)


def insert_paragraph(doc, text, align=WD_PARAGRAPH_ALIGNMENT.LEFT, size=9):
    """Insert a body paragraph with optional alignment and size."""
    para = doc.add_paragraph(text)
    para.alignment = align
    set_paragraph_font(para, size=size)
    return para


def insert_spacer(doc, size=6):
    """Insert an empty paragraph to create vertical spacing."""
    spacer = doc.add_paragraph()
    set_paragraph_font(spacer, size=size)
    return spacer


def insert_image_if_exists(doc, path, width=6.0, caption=None):
    """
    Insert an image if it exists, centered, with an optional caption.
    Width is expressed in inches.
    """
    if not os.path.exists(path):
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(path, width=Inches(width))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_font(paragraph, size=6)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_paragraph_font(cap, size=8)

    insert_spacer(doc, size=6)


def insert_table_from_csv(doc, csv_path, title=None):
    """
    Insert a table from a CSV (keeping column order).
    If the file does not exist or is empty, no action is taken.
    """
    if not os.path.exists(csv_path):
        return

    df = pd.read_csv(csv_path)
    if df.empty:
        return

    if title:
        insert_section_title(doc, title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # Headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)

    insert_spacer(doc, size=6)


# -------------------------------------------------------------------
# Report sections (1 … 8)
# -------------------------------------------------------------------

def _section_1_stations_context(doc, figures_dir, site_folder):
    """Stations context table (IDs, names, distances, coordinates)."""
    insert_section_title(doc, "Stations context", level=1)

    insert_paragraph(
        doc,
        (
            "This table provides information about the meteorological stations used "
            "(IDs, names, distances, coordinates)."
        ),
    )

    stations_csv_candidates = [
        os.path.join(figures_dir, "stations_context.csv"),
        os.path.join(site_folder, "stations_context.csv"),
    ]
    for path in stations_csv_candidates:
        if os.path.exists(path):
            insert_table_from_csv(doc, path)
            break


def _section_2_datas_quality(doc, figures_dir):
    """Data quality: study period and coverage by source."""
    insert_section_title(doc, "Data quality", level=1)

    insert_paragraph(
        doc,
        (
            "This table shows the number of days of available data, the study period, "
            "and the coverage rate for mean wind and gusts for each source used."
        ),
    )

    insert_table_from_csv(doc, os.path.join(figures_dir, "resume_qualite.csv"))


def _section_3_statistics_description(doc, figures_dir):
    """Descriptive statistics of mean wind speeds (m/s) by source."""
    insert_section_title(doc, "Statistics description", level=1)

    insert_paragraph(
        doc,
        (
            "Distribution of daily maximum mean wind speeds by source. For each "
            "dataset, the table shows the number of values, mean, standard deviation, "
            "percentiles and min/max. All speeds are expressed in m/s at 10 m height."
        ),
    )

    insert_table_from_csv(doc, os.path.join(figures_dir, "stats_windspeed_mean.csv"))


def _section_4_histograms(doc, figures_dir):
    """Histograms of mean wind speeds and gusts, by source."""
    insert_section_title(doc, "Histograms", level=1)

    insert_paragraph(
        doc,
        (
            "Histograms showing the distribution of daily maximum mean and gust wind "
            "speeds by source, with visualisation of the regulatory threshold "
            "(Building Code). These plots allow to visually compare the tails of "
            "each dataset."
        ),
    )

    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "histogrammes_windspeed_mean.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "histogrammes_windspeed_gust.png"),
        width=6.0,
    )


def _section_5_extreme_values(doc, figures_dir):
    """Extreme values analysis: boxplots, histograms of extreme days, summary tables."""
    insert_section_title(doc, "Extreme values analysis", level=1)

    insert_paragraph(
        doc,
        (
            "Visualisation of outliers via box plots and histograms for both mean "
            "wind and gusts. These extreme values may influence the analysis of "
            "mechanical loads and should be carefully checked in context (data "
            "quality, station history, possible measurement artefacts)."
        ),
    )

    # Boxplots + histograms of extreme days
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "boxplot_windspeed_mean.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "outliers_hist_windspeed_mean.png"),
        width=5.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "boxplot_windspeed_gust.png"),
        width=6.0,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "outliers_hist_windspeed_gust.png"),
        width=5.0,
    )

    # Extreme day tables - mean wind
    insert_paragraph(
        doc,
        "Summary of days above Building Code thresholds (mean wind, daily maxima).",
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "vent_moyen_extremes_resume.csv")
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "vent_moyen_extremes_par_an.csv")
    )

    # Extreme day tables - gusts
    insert_paragraph(
        doc,
        "Summary of days above Building Code thresholds (gust wind, daily maxima).",
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "rafales_extremes_resume.csv")
    )
    insert_table_from_csv(
        doc, os.path.join(figures_dir, "rafales_extremes_par_an.csv")
    )


def _section_6_directional_analysis(doc, figures_dir):
    """Directional analysis - compass roses (max speeds and occurrence)."""
    insert_section_title(doc, "Directional analysis - Compass roses", level=1)

    insert_paragraph(
        doc,
        (
            "Visualisation of prevailing wind directions by source. For each 20° "
            "sector between 0° and 360°, two roses are provided: "
            "one with the maximum daily mean wind speed by direction, and one with "
            "the frequency of occurrence. This helps to identify dominant wind "
            "sectors and their associated intensity."
        ),
    )

    if not os.path.isdir(figures_dir):
        return

    # Max speed roses
    for f in sorted(os.listdir(figures_dir)):
        if f.startswith("rose_max_windspeed_") and f.endswith(".png"):
            insert_image_if_exists(doc, os.path.join(figures_dir, f), width=5.5)

    # Frequency roses
    for f in sorted(os.listdir(figures_dir)):
        if f.startswith("rose_frequency_") and f.endswith(".png"):
            insert_image_if_exists(doc, os.path.join(figures_dir, f), width=5.5)


def _section_7_time_series(doc, figures_dir):
    """Full time series of daily maxima (mean wind / gust)."""
    insert_section_title(doc, "Time series", level=1)

    insert_paragraph(
        doc,
        (
            "View complete time series for daily maximum mean wind and gusts. "
            "These plots highlight long-term trends, potential regime changes and "
            "major extreme events (including outliers that may require a dedicated "
            "engineering judgement)."
        ),
    )

    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "time_series_windspeed_mean.png"),
        width=6.5,
    )
    insert_image_if_exists(
        doc,
        os.path.join(figures_dir, "time_series_windspeed_gust.png"),
        width=6.5,
    )


def _section_8_return_period_50y(doc, figures_dir):
    """Return levels (Gumbel) for 50 years, by source/variable."""
    insert_section_title(
        doc, "Analysis of levels for a 50-year return period", level=1
    )

    insert_paragraph(
        doc,
        (
            "Estimation of wind speed levels associated with a 50-year return "
            "period, using a Gumbel distribution fitted on daily maxima. "
            "Both mean wind and gusts are reported when enough data is available "
            "(typically at least 10 years of data). Results can be compared with "
            "Building Code thresholds to assess design margins."
        ),
    )

    insert_table_from_csv(doc, os.path.join(figures_dir, "return_period_50y.csv"))


# -------------------------------------------------------------------
# Main report generation
# -------------------------------------------------------------------

def generate_report(site_data, output_folder="data"):
    """
    Generate the Word report for a site, using the Word template (template.docx)
    for layout (logo, header/footer).
    Structure:
      1 | Stations context
      2 | Data quality
      3 | Statistics description
      4 | Histograms
      5 | Extreme values analysis
      6 | Directional analysis - Compass roses
      7 | Time series
      8 | Analysis of levels for a 50-year return period
    """
    site_name = f"{site_data['reference']}_{site_data['name']}"

    site_folder = os.path.join(output_folder, site_name)
    figures_dir = os.path.join(site_folder, "figures_and_tables")

    output_docx = os.path.join(site_folder, "report", f"{site_name}.docx")
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)

    # Load template if present
    template_path = os.path.join(os.getcwd(), "template.docx")
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        print("template.docx not found. Using a blank Word document.")
        doc = Document()

    # Intro banner
    insert_spacer(doc, size=6)
    intro = doc.add_paragraph()
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = intro.add_run(f"{site_data['reference']} - {site_data['name']}")
    run.bold = True
    run.font.size = Pt(16)

    insert_spacer(doc, size=10)

    # Sections 1 … 8
    _section_1_stations_context(doc, figures_dir, site_folder)
    doc.add_page_break()
    _section_2_datas_quality(doc, figures_dir)
    doc.add_page_break()
    _section_3_statistics_description(doc, figures_dir)
    doc.add_page_break()
    _section_4_histograms(doc, figures_dir)
    doc.add_page_break()
    _section_5_extreme_values(doc, figures_dir)
    doc.add_page_break()
    _section_6_directional_analysis(doc, figures_dir)
    doc.add_page_break()
    _section_7_time_series(doc, figures_dir)
    doc.add_page_break()
    _section_8_return_period_50y(doc, figures_dir)

    # Save
    doc.save(output_docx)
    print(f"DOCX report generated: {output_docx}")
